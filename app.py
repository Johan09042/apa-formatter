import os
import re
from flask import Flask, render_template, request, redirect, url_for, session, send_file
from docx import Document
import pdfplumber
from io import BytesIO
from openai import OpenAI

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "clave_segura")

# 🔐 LOGIN
USUARIO = os.getenv("USER", "admin")
PASSWORD = os.getenv("PASS", "1234")

# 🤖 IA SEGURA
client = None
api_key = os.getenv("OPENAI_API_KEY")
base_url = os.getenv("OPENAI_BASE_URL")

if api_key and base_url:
    try:
        client = OpenAI(api_key=api_key, base_url=base_url)
    except Exception as e:
        print("Error IA:", e)
else:
    print("⚠️ IA desactivada")

# 🔐 LOGIN
@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        user = request.form.get("usuario")
        password = request.form.get("password")

        if user == USUARIO and password == PASSWORD:
            session["login"] = True
            return redirect(url_for("index"))
        else:
            error = "Credenciales incorrectas"

    return render_template("login.html", error=error)

# 🔒 PROTEGER
@app.before_request
def proteger():
    rutas_libres = ["login", "static"]
    if request.endpoint not in rutas_libres:
        if not session.get("login"):
            return redirect(url_for("login"))

# 🚪 LOGOUT
@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# 🏠 HOME
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":

        archivo = request.files.get("archivo")
        usar_ia = request.form.get("usar_ia") == "on"
        usar_portada = request.form.get("usar_portada") == "on"
        instruccion = request.form.get("instruccion")

        tema = request.form.get("tema")
        nombre = request.form.get("nombre")
        materia = request.form.get("materia")
        profesor = request.form.get("profesor")
        colegio = request.form.get("colegio")

        texto = ""

        if archivo:
            if archivo.filename.endswith(".pdf"):
                with pdfplumber.open(archivo) as pdf:
                    for page in pdf.pages:
                        texto += page.extract_text() or ""
            elif archivo.filename.endswith(".docx"):
                doc = Document(archivo)
                for p in doc.paragraphs:
                    texto += p.text + "\n"
            else:
                texto = archivo.read().decode("utf-8", errors="ignore")

        contenido = texto

        # 🤖 IA COMPLETA
        if usar_ia and texto.strip() and client:
            try:
                prompt = f"""
Analiza el documento y responde EXACTAMENTE así:

---INTENCION---
...
---PARTES---
...
---ANTES---
...
---DESPUES---
...
---CAMBIOS---
...
---SUGERENCIAS---
...

Instrucciones del usuario: {instruccion}

Texto:
{texto}
"""
                response = client.chat.completions.create(
                    model="llama3-70b-8192",
                    messages=[
                        {"role": "system", "content": "Experto en redacción académica APA 7"},
                        {"role": "user", "content": prompt}
                    ]
                )
                contenido = response.choices[0].message.content
            except Exception as e:
                print("Error IA:", e)

        # 🔍 EXTRAER SECCIONES
        def extraer(nombre):
            try:
                patron = rf"---{nombre}---(.*?)(?=---|$)"
                match = re.search(patron, contenido, re.DOTALL)
                return match.group(1).strip() if match else "No disponible"
            except:
                return "No disponible"

        analisis = {
            "intencion": extraer("INTENCION"),
            "partes": extraer("PARTES"),
            "antes": extraer("ANTES"),
            "despues": extraer("DESPUES"),
            "cambios": extraer("CAMBIOS"),
            "sugerencias": extraer("SUGERENCIAS")
        }

        session["analisis"] = analisis
        session["portada"] = {
            "tema": tema,
            "nombre": nombre,
            "materia": materia,
            "profesor": profesor,
            "colegio": colegio,
            "usar": usar_portada
        }

        return redirect(url_for("preview"))

    return render_template("index.html")

# 📊 PREVIEW
@app.route("/preview")
def preview():
    analisis = session.get("analisis", {})
    return render_template("preview.html", analisis=analisis)

# 📄 DESCARGA WORD
@app.route("/download")
def download():
    analisis = session.get("analisis", {})
    portada = session.get("portada", {})

    doc = Document()

    # PORTADA
    if portada.get("usar"):
        doc.add_paragraph(portada.get("tema", ""))
        doc.add_paragraph(portada.get("nombre", ""))
        doc.add_paragraph(portada.get("materia", ""))
        doc.add_paragraph(portada.get("profesor", ""))
        doc.add_paragraph(portada.get("colegio", ""))
        doc.add_page_break()

    doc.add_paragraph(analisis.get("despues", ""))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="resultado_APA.docx")

# 🚀 PRODUCCIÓN
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
